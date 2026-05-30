"""Generate the MetaDrive eval summary Word document (data-driven, 3-arm).

Reads an eval_summary.json produced by rl/eval_metadrive.py, groups planners into
the three reward arms (Stock / Risk-only / Social-tuned) + IDM, and emits a
two-column (IEEE-style) Word doc with full-width summary tables and a discussion.

Usage:
  python docs/build_eval_summary_docx.py --summary rl/logs/metadrive/eval_intersection_3arm/eval_summary.json \
      --out docs/metadrive_intersection_eval_summary.docx
"""
import argparse, json, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ALGO_ORDER = ["ppo", "dqn", "sac", "td3", "ddpg"]
ARM_ORDER = ["Stock", "Risk-only", "Social-tuned", "IDM"]
ARM_SHADE = {"Stock": "FFFFFF", "Risk-only": "E2EFDA", "Social-tuned": "FCE4D6", "IDM": "EDEDED"}
HEADER_FILL = "1F4E79"


def classify(label: str):
    """label -> (arm, algo, display)."""
    l = label.lower()
    algo = next((a for a in ALGO_ORDER if l.endswith("_" + a) or ("_" + a + "_") in l), "ppo")
    if l.startswith("idm") or l == "idm":
        return "IDM", "idm", "IDM"
    if l.startswith("stock"):
        arm = "Stock"
    elif l.startswith("risk"):
        arm = "Risk-only"
    elif l.startswith("social") or "decoupled" in l or "lite" in l or "frozen" in l:
        arm = "Social-tuned"
    else:
        arm = "Stock"
    return arm, algo, f"{arm.split('-')[0]} {algo.upper()}"


def action_space(protocol: str) -> str:
    return "cont" if "continuous" in protocol else "disc"


def load_rows(summary_path: str):
    data = json.load(open(summary_path, encoding="utf-8"))
    rows = []
    for key, block in data.items():
        label, _, proto = key.partition("@")
        arm, algo, disp = classify(label)
        rows.append({"arm": arm, "algo": algo, "disp": disp,
                     "act": action_space(proto), "b": block})
    rows.sort(key=lambda r: (ARM_ORDER.index(r["arm"]) if r["arm"] in ARM_ORDER else 9,
                             ALGO_ORDER.index(r["algo"]) if r["algo"] in ALGO_ORDER else 9))
    return rows


def g(block, key, dp=2):
    m = block.get(key + "_mean")
    c = block.get(key + "_ci95")
    if m is None:
        return "—"
    return f"{m:.{dp}f}±{c:.{dp}f}" if c is not None else f"{m:.{dp}f}"


# ---- docx helpers ----
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:val'), 'clear'); s.set(qn('w:fill'), hexcolor); tcPr.append(s)

def set_cell(cell, text, *, bold=False, size=8, white=False, fill=None, align='center'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.name = 'Arial'
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill: shade(cell, fill)

def fixed_layout(tbl):
    tbl.autofit = False; tbl.allow_autofit = False
    el = OxmlElement('w:tblLayout'); el.set(qn('w:type'), 'fixed'); tbl._tbl.tblPr.append(el)

def set_cols(section, num, space=480):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), str(space))

def build_table(doc, rows, headers, specs, widths):
    """specs: list of (metric_key, dp); widths sum ~6.5in; col0=label, col1=Act."""
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = 'Table Grid'; fixed_layout(tbl)
    for c, h, w in zip(tbl.rows[0].cells, headers, widths):
        set_cell(c, h, bold=True, white=True, fill=HEADER_FILL); c.width = Inches(w)
    for r in rows:
        cells = tbl.add_row().cells
        fill = ARM_SHADE.get(r["arm"], "FFFFFF")
        set_cell(cells[0], r["disp"], bold=(r["arm"] != "Stock"), fill=fill, align='left'); cells[0].width = Inches(widths[0])
        set_cell(cells[1], r["act"], fill=fill); cells[1].width = Inches(widths[1])
        for j, (k, dp) in enumerate(specs, start=2):
            set_cell(cells[j], g(r["b"], k, dp), fill=fill); cells[j].width = Inches(widths[j])

def add_caption(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(8); r.font.name = 'Arial'

def add_para(doc, lead, body):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r1 = p.add_run(lead + " "); r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = 'Arial'
    r2 = p.add_run(body); r2.font.size = Pt(9.5); r2.font.name = 'Arial'


# Discussion paragraphs — finalized after the 20-seed numbers land. Edit DISCUSSION
# and rerun. (lead, body) tuples; lead is bolded.
DISCUSSION = [
    ("Experimental setup.",
     "Three reward arms — stock (r_MD), risk-only (r_MD minus a DRIFT risk penalty, with the 8-D risk "
     "observation but no comfort/courtesy terms), and social-tuned (risk-only plus the decoupled, calibrated "
     "ego-comfort and social-externality shaping) — are compared across five RL families on the "
     "unprotected-intersection map with respawning traffic (density 0.3), over 20 seed-paired episodes; "
     "values are mean ± 95% CI. The clean per-algorithm comparison is the off-policy continuous trio "
     "SAC/TD3/DDPG (all three arms share the continuous action space); PPO mixes action spaces across arms "
     "and DQN is discrete-only, as flagged in the Act column."),
    ("Off-policy continuous control is required.",
     "PPO and DQN fail on this scenario in every arm (0% success, route ≤ 0.39, and the continuous "
     "social-PPO even degenerates to 90% crash). Only the off-policy continuous learners SAC/TD3/DDPG reach "
     "0.45–0.70 success, so the arm comparison is drawn over that trio. Algorithm family dominates the "
     "reward regime in determining competence on this geometry."),
    ("The DRIFT risk field is a clean win for the strongest learner.",
     "Risk-only SAC is the best single policy in the study: success 0.70 (stock 0.60), route 0.770 (0.686), "
     "crash 0.20 (0.30), follower-deceleration 0.025 (0.034), and the highest social-friendliness composite "
     "of all fifteen policies (0.787). Adding the risk observation and penalty improves safety and progress "
     "simultaneously at no comfort cost. The effect is algorithm-dependent — smaller for DDPG and slightly "
     "negative for TD3 (success 0.50 vs 0.60) — indicating the propagated-risk signal helps the most "
     "sample-efficient learner most."),
    ("Full social shaping yields algorithm-dependent comfort/courtesy gains.",
     "The tuned social reward lowers jerk for SAC (15.1 vs 15.5) and DDPG (17.6 vs 18.7) and gives SAC the "
     "lowest follower-deceleration in the study (0.022). Social DDPG is the strongest DDPG variant overall "
     "(success 0.65, lowest crash 0.20, highest DDPG composite 0.767). The benefit is not universal: social "
     "TD3 worsens follower-deceleration (0.052) and on-policy social PPO degenerates, so the shaping must be "
     "applied per-algorithm rather than uniformly."),
    ("Comfort shaping trades against lane-keeping for some algorithms.",
     "Social SAC attains its smoothness/courtesy gains at a higher off-road rate (0.30 vs 0.15 for stock), "
     "which at 20 seeds is outside noise: the residual jerk penalty suppresses the sharp corrective steering "
     "needed at the intersection — the same mechanism identified in the weight-calibration study. A modest "
     "reduction of the jerk weight (toward 0.001) is the indicated remedy before extending to other maps."),
    ("Risk-awareness alone reduces follower disturbance.",
     "Across the off-policy trio the risk-only arm yields the lowest follower-deceleration (SAC 0.025, "
     "TD3 0.016, DDPG 0.023, each below its stock counterpart), suggesting the propagated risk field "
     "discourages aggressive gap-closing even without an explicit courtesy term; the courtesy/backward-flux "
     "terms then sharpen this effect for SAC."),
    ("Deployability.",
     "All learned policies run at ~1–3 ms per step on CPU and IDM is effectively free, comfortably real-time "
     "across families; the 134 ms TD3 figure seen in an earlier run was measurement contention and does not "
     "reproduce here."),
    ("Takeaway.",
     "The DRIFT risk field is a reliable, low-cost improvement to the strongest RL driver (risk-SAC dominates "
     "the table), while the added comfort/courtesy shaping delivers genuine smoothness and courtesy for SAC "
     "and DDPG when tuned, at a route-completion trade-off that is itself controllable through the jerk "
     "weight. PPO and DQN remain uncompetitive on this geometry regardless of reward arm."),
]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--summary", required=True)
    ap.add_argument("--out", default="docs/metadrive_intersection_eval_summary.docx")
    args = ap.parse_args()
    rows = load_rows(args.summary)

    doc = Document()
    s0 = doc.sections[0]
    s0.page_width = Inches(8.5); s0.page_height = Inches(11)
    s0.top_margin = s0.bottom_margin = s0.left_margin = s0.right_margin = Inches(1)
    set_cols(s0, 1)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Closed-Loop 3-Arm Ablation in MetaDrive: Stock vs. Risk-Only vs. Social-Tuned RL")
    tr.bold = True; tr.font.size = Pt(14); tr.font.name = 'Arial'
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("Unprotected Intersection, Respawn Traffic — 20 paired seeds, density 0.3 (mean ± 95% CI)")
    sr.font.size = Pt(10); sr.italic = True; sr.font.name = 'Arial'

    W1 = [1.30, 0.55, 0.93, 0.93, 0.93, 0.93, 0.93]  # 7 cols ~6.5"
    add_caption(doc, "Table 1. Safety and progress by reward arm and algorithm.")
    build_table(doc, rows,
        ["Planner", "Act", "Success", "Crash", "Off-road", "Route", "Base reward"],
        [("success", 2), ("crash_any", 2), ("out_of_road", 2), ("route_completion", 3), ("ep_base_reward", 1)],
        W1)

    W2 = [1.30, 0.55, 0.78, 0.84, 0.86, 0.90, 0.74, 0.53]  # 8 cols ~6.5"
    add_caption(doc, "Table 2. Comfort, social-externality, and inference cost by reward arm and algorithm.")
    build_table(doc, rows,
        ["Planner", "Act", "Jerk", "SteerΔ", "Comfort", "FollDecel", "Social", "ms"],
        [("mean_jerk_abs", 1), ("steering_change_rate", 2), ("mean_comfort_cost", 3),
         ("mean_decel_follower", 3), ("social_friendliness_score", 3), ("mean_action_selection_ms", 1)],
        W2)

    note = doc.add_paragraph(); note.paragraph_format.space_after = Pt(8)
    nr = note.add_run("Near-miss/TTC counts ≈0 and peak risk ≈0.11 across planners (omitted). 'Act' = action "
                      "space (cont/disc). Social-tuned weights: λ_R=0.02, λ_jerk=0.0015, λ_Δδ=0.001, "
                      "w_courtesy=0.05, w_rear_ttc=0.02, w_back_flux=0.02.")
    nr.font.size = Pt(8); nr.italic = True; nr.font.name = 'Arial'

    doc.add_section(WD_SECTION.CONTINUOUS)
    s1 = doc.sections[-1]
    s1.page_width = Inches(8.5); s1.page_height = Inches(11)
    s1.top_margin = s1.bottom_margin = s1.left_margin = s1.right_margin = Inches(1)
    set_cols(s1, 2, space=480)
    h = doc.add_paragraph(); hr = h.add_run("Analysis"); hr.bold = True; hr.font.size = Pt(12); hr.font.name = 'Arial'
    h.paragraph_format.space_after = Pt(4)
    for lead, body in DISCUSSION:
        add_para(doc, lead, body)

    doc.save(args.out)
    print("saved", args.out)
    d2 = Document(args.out)
    print("rows:", len(rows), "tables:", len(d2.tables), "t1rows:", len(d2.tables[0].rows), "sections:", len(d2.sections))
    # echo key numbers for writing the discussion
    print("\nkey numbers:")
    for r in rows:
        b = r["b"]
        print(f"  {r['disp']:14s}[{r['act']}] succ={g(b,'success')} crash={g(b,'crash_any')} oor={g(b,'out_of_road')} "
              f"route={g(b,'route_completion',3)} jerk={g(b,'mean_jerk_abs',1)} follDcl={g(b,'mean_decel_follower',3)} "
              f"social={g(b,'social_friendliness_score',3)} ms={g(b,'mean_action_selection_ms',1)}")


if __name__ == "__main__":
    main()
