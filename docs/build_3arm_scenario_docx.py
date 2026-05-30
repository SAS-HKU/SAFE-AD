"""Generate per-scenario 3-arm Word docs (intersection/merge/roundabout) from the
eval_*_3arm_full summaries. Reuses table helpers from build_eval_summary_docx.py
(non-destructive: new filenames metadrive_<scenario>_3arm_summary.docx).
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from build_eval_summary_docx import load_rows, build_table, add_caption, add_para, set_cols

SUMMARY = "rl/logs/metadrive/eval_{S}_3arm_full/eval_summary.json"
OUT = "docs/metadrive_{S}_3arm_summary.docx"
TITLES = {"intersection": "Unprotected Intersection", "merge": "On-Ramp Merge", "roundabout": "Roundabout"}

SETUP = ("Three reward arms — stock (r_MD), risk-only (r_MD with the 8-D DRIFT risk observation and a risk "
         "penalty but no comfort/courtesy terms), and social-tuned (risk-only plus the decoupled, calibrated "
         "ego-comfort and social-externality shaping) — are compared across five RL families and the native "
         "IDM planner over 20 seed-paired episodes at traffic density 0.3; values are mean ± 95% CI. The clean "
         "per-algorithm comparison is the off-policy continuous trio SAC/TD3/DDPG (all arms continuous); PPO "
         "mixes action spaces across arms and DQN is discrete-only (Act column).")

CROSS = ("Across all three scenarios, SAC is the only RL family that drives competently everywhere; TD3 "
         "collapses to a near-static policy on merge and roundabout (≈0% success, route ≈0.03, jerk ≈7.8) and "
         "DDPG is inconsistent. The DRIFT risk field (risk-only) gives SAC a consistent gain over stock "
         "(+0.10 success on intersection, +0.15 on merge) or a tie, and the full social-tuned reward is "
         "strongest on the roundabout. The social benefit is therefore scenario- and algorithm-dependent and "
         "should be read over the SAC backbone.")

DISC = {
 "intersection": [
  ("Off-policy continuous control is required.",
   "PPO and DQN fail in every arm (0% success, route ≤ 0.39); only SAC/TD3/DDPG are competent, so the arm "
   "comparison is drawn over them."),
  ("The DRIFT risk field is a clean win for the strongest learner.",
   "Risk-only SAC is the best policy here: success 0.70 (stock 0.60), route 0.770 (0.692), crash 0.20 (0.35), "
   "follower-deceleration 0.025 (0.049), and the highest social composite (0.787). Adding the risk channel "
   "improves safety and progress at once."),
  ("Full social shaping is smoothest and most courteous, with an off-road cost for SAC.",
   "Social SAC attains the lowest jerk (15.1) and follower-deceleration (0.022) in the scenario but raises "
   "off-road to 0.30, lowering success to 0.55. Social DDPG is the strongest DDPG variant (success 0.60, "
   "crash 0.20, composite 0.764), showing the shaping helps when it does not destabilise lane-keeping."),
 ],
 "merge": [
  ("Off-policy fragility dominates the merge.",
   "TD3 collapses in all arms to a near-static policy (success 0, route 0.026, jerk 7.8), as do risk- and "
   "social-DDPG; only SAC (and stock-DDPG) drive. The limiting factor on merge is optimizer stability, not "
   "the reward."),
  ("Risk and social shaping both lift the stable learner.",
   "Risk-SAC and social-SAC tie at the top — success 0.80 and route 0.806 versus stock SAC's 0.65 / 0.688 — "
   "and social-SAC is smoother (jerk 15.0 vs 17.1) with zero imposed follower braking. Both reach a 0.770 "
   "social composite."),
  ("IDM remains the strongest merge planner.",
   "The rule-based IDM leads on success (0.85) and composite (0.809); the learned social-SAC is the closest "
   "RL competitor and the smoothest among the high-success policies."),
 ],
 "roundabout": [
  ("Full social shaping wins outright on the roundabout.",
   "Social SAC is the best of all sixteen policies: success 0.85, crash 0.15 (lowest), route 0.835, and the "
   "top social composite (0.787) — beating stock SAC (0.70), risk SAC (0.70), and IDM (0.65). Here the social "
   "reward improves safety and progress simultaneously, with no off-road penalty."),
  ("Stock DDPG is the other strong policy; TD3 again collapses.",
   "Stock DDPG reaches success 0.80 / route 0.820, but its risk and social variants regress; TD3 collapses in "
   "every arm (success 0). The roundabout geometry rewards the smooth, courteous SAC policy the social "
   "shaping produces."),
 ],
}


def build_doc(scenario):
    rows = load_rows(SUMMARY.format(S=scenario))
    doc = Document()
    s0 = doc.sections[0]
    s0.page_width = Inches(8.5); s0.page_height = Inches(11)
    s0.top_margin = s0.bottom_margin = s0.left_margin = s0.right_margin = Inches(1)
    set_cols(s0, 1)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(f"MetaDrive 3-Arm Ablation — {TITLES[scenario]}")
    tr.bold = True; tr.font.size = Pt(14); tr.font.name = 'Arial'
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("Stock vs. Risk-only vs. Social-tuned RL — 20 paired seeds, respawn traffic, density 0.3 (mean ± 95% CI)")
    sr.font.size = Pt(10); sr.italic = True; sr.font.name = 'Arial'

    W1 = [1.30, 0.55, 0.93, 0.93, 0.93, 0.93, 0.93]
    add_caption(doc, "Table 1. Safety and progress by reward arm and algorithm.")
    build_table(doc, rows,
        ["Planner", "Act", "Success", "Crash", "Off-road", "Route", "Base reward"],
        [("success", 2), ("crash_any", 2), ("out_of_road", 2), ("route_completion", 3), ("ep_base_reward", 1)], W1)
    W2 = [1.30, 0.55, 0.78, 0.84, 0.86, 0.90, 0.74, 0.53]
    add_caption(doc, "Table 2. Comfort, social-externality, and inference cost by reward arm and algorithm.")
    build_table(doc, rows,
        ["Planner", "Act", "Jerk", "SteerΔ", "Comfort", "FollDecel", "Social", "ms"],
        [("mean_jerk_abs", 1), ("steering_change_rate", 2), ("mean_comfort_cost", 3),
         ("mean_decel_follower", 3), ("social_friendliness_score", 3), ("mean_action_selection_ms", 1)], W2)

    note = doc.add_paragraph(); note.paragraph_format.space_after = Pt(8)
    nr = note.add_run("Near-miss/TTC counts ≈0 and peak risk ≈0.11 across planners (omitted). Social-tuned weights: "
                      "λ_R=0.02, λ_jerk=0.0015, λ_Δδ=0.001, w_courtesy=0.05, w_rear_ttc=0.02, w_back_flux=0.02.")
    nr.font.size = Pt(8); nr.italic = True; nr.font.name = 'Arial'

    doc.add_section(WD_SECTION.CONTINUOUS)
    s1 = doc.sections[-1]
    s1.page_width = Inches(8.5); s1.page_height = Inches(11)
    s1.top_margin = s1.bottom_margin = s1.left_margin = s1.right_margin = Inches(1)
    set_cols(s1, 2, space=480)
    h = doc.add_paragraph(); hr = h.add_run("Analysis"); hr.bold = True; hr.font.size = Pt(12); hr.font.name = 'Arial'
    add_para(doc, "Experimental setup.", SETUP)
    for lead, body in DISC[scenario]:
        add_para(doc, lead, body)
    add_para(doc, "Cross-scenario context.", CROSS)
    out = OUT.format(S=scenario); doc.save(out)
    return out, len(rows), len(doc.tables)


if __name__ == "__main__":
    for sc in ["intersection", "merge", "roundabout"]:
        out, nr, nt = build_doc(sc)
        print(f"saved {out}  ({nr} planners, {nt} tables)")
