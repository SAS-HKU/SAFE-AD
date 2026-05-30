"""Single combined cross-scenario 3-arm Word doc (intersection+merge+roundabout).

Top: title, setup, a compact SAC headline table, then per-scenario full tables.
Bottom (2-col): per-scenario discussions + cross-scenario summary.
Reuses helpers from build_eval_summary_docx.py and discussions from
build_3arm_scenario_docx.py. Non-destructive (new file).
"""
import os, sys, json
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from build_eval_summary_docx import (load_rows, build_table, add_caption, add_para,
                                     set_cols, set_cell, fixed_layout, g, HEADER_FILL)
from build_3arm_scenario_docx import DISC, SETUP, CROSS, TITLES, SUMMARY

OUT = "docs/metadrive_3arm_combined_summary.docx"
SCEN = ["intersection", "merge", "roundabout"]

T1 = (["Planner", "Act", "Success", "Crash", "Off-road", "Route", "Base reward"],
      [("success", 2), ("crash_any", 2), ("out_of_road", 2), ("route_completion", 3), ("ep_base_reward", 1)],
      [1.30, 0.55, 0.93, 0.93, 0.93, 0.93, 0.93])
T2 = (["Planner", "Act", "Jerk", "SteerΔ", "Comfort", "FollDecel", "Social", "ms"],
      [("mean_jerk_abs", 1), ("steering_change_rate", 2), ("mean_comfort_cost", 3),
       ("mean_decel_follower", 3), ("social_friendliness_score", 3), ("mean_action_selection_ms", 1)],
      [1.30, 0.55, 0.78, 0.84, 0.86, 0.90, 0.74, 0.53])


def heading(doc, text, size=12):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.name = 'Arial'


def headline_table(doc):
    """Compact SAC success/social-composite across 3 arms x 3 scenarios."""
    sums = {S: {k.split("@")[0]: v for k, v in json.load(open(SUMMARY.format(S=S))).items()} for S in SCEN}
    headers = ["Scenario", "Stock SAC", "Risk SAC", "Social SAC", "IDM"]
    widths = [1.40, 1.28, 1.28, 1.28, 1.26]
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = 'Table Grid'; fixed_layout(tbl)
    for c, h, w in zip(tbl.rows[0].cells, headers, widths):
        set_cell(c, h, bold=True, white=True, fill=HEADER_FILL); c.width = Inches(w)
    for S in SCEN:
        by = sums[S]; cells = tbl.add_row().cells
        set_cell(cells[0], TITLES[S], bold=True, align='left'); cells[0].width = Inches(widths[0])
        for j, lab in enumerate(["stock_sac", "risk_sac", "social_sac", "idm"], start=1):
            b = by.get(lab, {})
            val = f"{g(b,'success',2)} / {g(b,'social_friendliness_score',2)}" if b else "—"
            set_cell(cells[j], val); cells[j].width = Inches(widths[j])


def main():
    doc = Document()
    s0 = doc.sections[0]
    s0.page_width = Inches(8.5); s0.page_height = Inches(11)
    s0.top_margin = s0.bottom_margin = s0.left_margin = s0.right_margin = Inches(1)
    set_cols(s0, 1)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("MetaDrive Closed-Loop 3-Arm Ablation across Three Scenarios")
    tr.bold = True; tr.font.size = Pt(15); tr.font.name = 'Arial'
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("Stock vs. Risk-only vs. Social-tuned RL — 5 algorithms × 3 scenarios, "
                    "20 paired seeds, respawn traffic, density 0.3 (mean ± 95% CI)")
    sr.font.size = Pt(10); sr.italic = True; sr.font.name = 'Arial'

    add_caption(doc, "Table 1. Headline: SAC success / social-friendliness composite by arm and scenario.")
    headline_table(doc)

    for i, S in enumerate(SCEN, start=2):
        rows = load_rows(SUMMARY.format(S=S))
        heading(doc, f"{TITLES[S]}")
        add_caption(doc, f"Table {2*i-2}. {TITLES[S]}: safety and progress.")
        build_table(doc, rows, *T1)
        add_caption(doc, f"Table {2*i-1}. {TITLES[S]}: comfort, social-externality, and inference cost.")
        build_table(doc, rows, *T2)

    note = doc.add_paragraph(); note.paragraph_format.space_before = Pt(6); note.paragraph_format.space_after = Pt(6)
    nr = note.add_run("Act = action space (cont/disc). Near-miss/TTC ≈0 and peak risk ≈0.11 across planners "
                      "(omitted). Social-tuned weights: λ_R=0.02, λ_jerk=0.0015, λ_Δδ=0.001, w_courtesy=0.05, "
                      "w_rear_ttc=0.02, w_back_flux=0.02.")
    nr.font.size = Pt(8); nr.italic = True; nr.font.name = 'Arial'

    # ---- embedded results figures (full width, 1-col) ----
    figs = [
        ("rl/logs/figures/social_sac_reward_return.png",
         "Figure 1. Social-SAC training dynamics — average reward and episode return across the three "
         "scenarios (identical objective; scienceplots styling)."),
        ("rl/logs/figures/social_sac_decomp_training.png",
         "Figure 2. Social-SAC reward-term decomposition over training — per-term penalty magnitude by "
         "scenario; risk and courtesy terms dominate early and decay as the policy learns."),
    ]
    have_fig = False
    for path, cap in figs:
        if os.path.exists(path):
            have_fig = True
            doc.add_picture(path, width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, cap)
    if have_fig:
        # keep the figures on the tables page; start discussion fresh below
        pass

    # ---- 2-column discussion ----
    doc.add_section(WD_SECTION.CONTINUOUS)
    s1 = doc.sections[-1]
    s1.page_width = Inches(8.5); s1.page_height = Inches(11)
    s1.top_margin = s1.bottom_margin = s1.left_margin = s1.right_margin = Inches(1)
    set_cols(s1, 2, space=480)
    h = doc.add_paragraph(); hr = h.add_run("Analysis"); hr.bold = True; hr.font.size = Pt(12); hr.font.name = 'Arial'
    add_para(doc, "Experimental setup.", SETUP)
    for S in SCEN:
        add_para(doc, f"{TITLES[S]}.", " ".join(b for _, b in DISC[S]))
    add_para(doc, "Off-road tuning is a genuine trade-off.",
             "On the intersection, social SAC's one weakness is off-road (0.30). Lowering the jerk weight "
             "(λ_jerk 0.0015→0.001) reduces it to 0.20 but raises crash (0.20→0.30) and triples "
             "follower-deceleration (0.022→0.060), lowering the social composite (0.760→0.735); most rate "
             "differences lie within the 20-seed CIs. We therefore retain the decoupled weights as the "
             "social-friendly operating point and report off-road as a documented limitation.")
    add_para(doc, "Cross-scenario summary.", CROSS)

    out = OUT
    try:
        doc.save(out)
    except PermissionError:
        out = OUT.replace(".docx", "_withfigs.docx")
        doc.save(out)
        print(f"[warn] {OUT} is locked (open in Word?); saved to {out} instead")
    d2 = Document(out)
    print(f"saved {out}  tables={len(d2.tables)} sections={len(d2.sections)} paras={len(d2.paragraphs)}")


if __name__ == "__main__":
    main()
